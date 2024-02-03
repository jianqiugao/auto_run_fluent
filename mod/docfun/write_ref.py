import os.path

import docx
from docx.shared import RGBColor
from docx.oxml.ns import qn  # 中文格式
from docx.shared import Inches, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


def write_ref(document, intial_counts, parents_path):
    size_0 = 16
    size_1 = 10
    size_2 = 8
    first = document.add_heading("一、参考依据", level=2)
    intial_counts = intial_counts + 1
    run = first.runs[0]
    run.font.color.rgb = RGBColor(0, 0, 0)  # 设置文本颜色为红色
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    run.font.size = docx.shared.Pt(size_0)  # 设置第一个运行的字体大小为15磅

    first_first = document.add_heading("1.技术需求", level=3)
    intial_counts = intial_counts + 1
    run = first_first.runs[0]
    run.font.color.rgb = RGBColor(0, 0, 0)  # 设置文本颜色为红色
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    run.font.size = docx.shared.Pt(size_1)  # 设置第一个运行的字体大小为15磅

    picture = document.add_picture(os.path.join(parents_path, 'file/basepic/tech_requirement.png'),
                                   width=Inches(6.25))  # 设置图片宽度，inches（英尺）与cm（厘米）两种
    document.paragraphs[intial_counts].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 居中对齐
    intial_counts = intial_counts + 1

    first_second = document.add_heading("2.合同约定", level=3)
    run = first_second.runs[0]
    run.font.color.rgb = RGBColor(0, 0, 0)  # 设置文本颜色为红色
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    run.font.size = docx.shared.Pt(size_1)  # 设置第一个运行的字体大小为15磅
    intial_counts = intial_counts + 1

    first_second_first = document.add_paragraph("1).软件接口方面,基于现有商业软件ansys fluent,Ensight开发;")
    run = first_second_first.runs[0]
    run.font.color.rgb = RGBColor(0, 0, 0)  # 设置文本颜色为红色
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    run.font.size = docx.shared.Pt(size_2)  # 设置第一个运行的字体大小为15磅
    intial_counts = intial_counts + 1

    first_second_second = document.add_paragraph("2).实现作业计算自动监控，并在分析完成后自动进行结果分析输出;")
    run = first_second_second.runs[0]
    run.font.color.rgb = RGBColor(0, 0, 0)  # 设置文本颜色为红色
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    run.font.size = docx.shared.Pt(size_2)  # 设置第一个运行的字体大小为15磅
    intial_counts = intial_counts + 1

    first_second_third = document.add_paragraph("3).实现分析报告自动输出,报告格式为word/pdf/ppt;")
    run = first_second_third.runs[0]
    run.font.color.rgb = RGBColor(0, 0, 0)  # 设置文本颜色为红色
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    run.font.size = docx.shared.Pt(size_2)  # 设置第一个运行的字体大小为15磅
    intial_counts = intial_counts + 1

    first_second_third = document.add_paragraph("4).报告内容要求包括但不限于仿真需求中设置的工况条件,模型设置,数据曲线对比,云图等;")
    run = first_second_third.runs[0]
    run.font.color.rgb = RGBColor(0, 0, 0)  # 设置文本颜色为红色
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    run.font.size = docx.shared.Pt(8)  # 设置第一个运行的字体大小为15磅
    intial_counts = intial_counts + 1

    first_second_third = document.add_paragraph("5).开发语言采用底层计算机语言，适用于Windows操作系统;")
    run = first_second_third.runs[0]
    run.font.color.rgb = RGBColor(0, 0, 0)  # 设置文本颜色为红色
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    run.font.size = docx.shared.Pt(8)  # 设置第一个运行的字体大小为15磅
    intial_counts = intial_counts + 1

    first_third = document.add_heading("3.相关标准文件", level=3)
    run = first_third.runs[0]
    run.font.color.rgb = RGBColor(0, 0, 0)  # 设置文本颜色为红色
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    run.font.size = docx.shared.Pt(size_1)  # 设置第一个运行的字体大小为15磅
    intial_counts = intial_counts + 1

    first_third_first = document.add_paragraph("仿真结果分析及可视化应符合以下要求")
    run = first_third_first.runs[0]
    run.font.color.rgb = RGBColor(0, 0, 0)  # 设置文本颜色为红色
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    run.font.size = docx.shared.Pt(size_2)  # 设置第一个运行的字体大小为15磅
    intial_counts = intial_counts + 1

    first_third_second = document.add_paragraph("a) 系统仿真应能够进行工时、材料定额的计算和统计，辅助进行材料及人工成本分析")
    run = first_third_second.runs[0]
    run.font.color.rgb = RGBColor(0, 0, 0)  # 设置文本颜色为红色
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    run.font.size = docx.shared.Pt(size_2)  # 设置第一个运行的字体大小为15磅
    intial_counts = intial_counts + 1

    first_third_third = document.add_paragraph("b) 系统仿真应根据不同的系统特点、仿真要求对仿真数据进行取舍分析，得到评价结果;")
    run = first_third_third.runs[0]
    run.font.color.rgb = RGBColor(0, 0, 0)  # 设置文本颜色为红色
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    run.font.size = docx.shared.Pt(size_2)  # 设置第一个运行的字体大小为15磅
    intial_counts = intial_counts + 1

    first_third_fourth = document.add_paragraph(
        "c) 应能通过汇总分析加工精度、设备利用率、加工方法、加工效率等信息，综合判断NC程序是否合理，辅助进行NC程序优化，并能对优化前后评价因素进行比对，分析或辅助分析改进情况;")
    run = first_third_fourth.runs[0]
    run.font.color.rgb = RGBColor(0, 0, 0)  # 设置文本颜色为红色
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    run.font.size = docx.shared.Pt(size_2)  # 设置第一个运行的字体大小为15磅
    intial_counts = intial_counts + 1

    first_third_fi = document.add_paragraph(
        "d) 对于人体模型应能对工效分析结果进行分析，辅助进行人体操作改进，在进行优化改进时，应d能从姿态分析、视野分析、作业空间、工效分析等方面对比优化改进情况;")
    run = first_third_fi.runs[0]
    run.font.color.rgb = RGBColor(0, 0, 0)  # 设置文本颜色为红色
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    run.font.size = docx.shared.Pt(size_2)  # 设置第一个运行的字体大小为15磅
    intial_counts = intial_counts + 1

    first_third_six = document.add_paragraph("e) 生产线系统仿真结果评价应包含生产线瓶颈识别和分析，包括产生瓶颈的设备位置和程度，其指标有生产线的瓶颈率、等待队长、等待时间等:")
    run = first_third_six.runs[0]
    run.font.color.rgb = RGBColor(0, 0, 0)  # 设置文本颜色为红色
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    run.font.size = docx.shared.Pt(size_2)  # 设置第一个运行的字体大小为15磅
    intial_counts = intial_counts + 1

    first_third_seven = document.add_paragraph("f) 生产线系统仿真结果评价应包含车间产能分析，包括最大能力和负载率，其指标包括单位时间生产量、初始加工时间、设备利用率等;")
    run = first_third_seven.runs[0]
    run.font.color.rgb = RGBColor(0, 0, 0)  # 设置文本颜色为红色
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    run.font.size = docx.shared.Pt(size_2)  # 设置第一个运行的字体大小为15磅
    intial_counts = intial_counts + 1

    first_third_seven = document.add_paragraph("g) 生产线系统仿真结果评价还应包含生产周期分析以及生产的人力需求分析;")
    run = first_third_seven.runs[0]
    run.font.color.rgb = RGBColor(0, 0, 0)  # 设置文本颜色为红色
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    run.font.size = docx.shared.Pt(size_2)  # 设置第一个运行的字体大小为15磅
    intial_counts = intial_counts + 1
    return intial_counts
