from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import RGBColor
from docx.oxml.ns import qn  # 中文格式
import docx
import datetime


def write_summary(document, run_dir,init_numm,promts):
    size_0 = 16
    size_1 = 10
    second = document.add_heading("四、报告总结", level=2)
    run = second.runs[0]
    run.font.color.rgb = RGBColor(0, 0, 0)  # 设置文本颜色为红色
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    run.font.size = docx.shared.Pt(size_0)  # 设置第一个运行的字体大小为15磅
    init_numm = init_numm+1
    title = document.add_paragraph(f"报告名称:{run_dir}")
    run = title.runs[0]
    run.font.color.rgb = RGBColor(0, 0, 0)  # 设置文本颜色为红色

    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    run.font.size = docx.shared.Pt(size_1)  # 设置第一个运行的字体大小为15磅
    init_numm = init_numm + 1
    sumary = promts['promts']  # todo 增加校验函数
    # 计算结论
    title = document.add_paragraph(f"计算结论:{sumary}")
    run = title.runs[0]
    run.font.color.rgb = RGBColor(0, 0, 0)  # 设置文本颜色为红色
    # title.text = "仿真分析报告"
    # 设置西体格式
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    run.font.size = docx.shared.Pt(10)  # 设置第一个运行的字体大小为15磅
    init_numm = init_numm + 1

    # 编写
    editor = promts['editor']#'computer auto writer'
    title = document.add_paragraph(f"编写:{editor}")
    run = title.runs[0]
    run.font.color.rgb = RGBColor(0, 0, 0)  # 设置文本颜色为红色
    # title.text = "仿真分析报告"
    # 设置西体格式
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    run.font.size = docx.shared.Pt(10)  # 设置第一个运行的字体大小为15磅
    init_numm = init_numm + 1

    # 审核
    checker = promts['checker']
    title = document.add_paragraph(f"审核:{checker}")
    run = title.runs[0]
    run.font.color.rgb = RGBColor(0, 0, 0)  # 设置文本颜色为红色
    # title.text = "仿真分析报告"
    # 设置西体格式
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    run.font.size = docx.shared.Pt(size_1)  # 设置第一个运行的字体大小为15磅
    init_numm = init_numm + 1

    # 审核
    author = promts['auth']
    title = document.add_paragraph(f"批准:{author}")
    run = title.runs[0]
    run.font.color.rgb = RGBColor(0, 0, 0)  # 设置文本颜色为红色
    # title.text = "仿真分析报告"
    # 设置西体格式
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    run.font.size = docx.shared.Pt(size_1)  # 设置第一个运行的字体大小为15磅
    # 审核
    data = datetime.datetime.now().strftime(promts['DateFormate'])
    title = document.add_paragraph(f"日期:{data}")
    run = title.runs[0]
    run.font.color.rgb = RGBColor(0, 0, 0)  # 设置文本颜色为红色
    # title.text = "仿真分析报告"
    # 设置西体格式
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    run.font.size = docx.shared.Pt(size_1)  # 设置第一个运行的字体大小为15磅

    pass
