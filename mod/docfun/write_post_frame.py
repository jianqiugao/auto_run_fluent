import os.path

import docx
from docx.shared import RGBColor
from docx.oxml.ns import qn  # 中文格式
from docx.shared import Inches, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_ALIGN_PARAGRAPH


def write_post_frame(document, intial_counts, work_condition_params, order_params, parent_path, run_dir):
    size_0 = 16
    size_1 = 10
    size_2 = 8
    second = document.add_heading("二、后处理框架内容", level=2)
    run = second.runs[0]
    run.font.color.rgb = RGBColor(0, 0, 0)  # 设置文本颜色为红色
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    run.font.size = docx.shared.Pt(size_0)  # 设置第一个运行的字体大小为15磅
    intial_counts = intial_counts + 1

    second_first = document.add_heading("1.输入参数", level=3)
    run = second_first.runs[0]
    run.font.color.rgb = RGBColor(0, 0, 0)  # 设置文本颜色为红色
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    run.font.size = docx.shared.Pt(size_1)  # 设置第一个运行的字体大小为15磅
    intial_counts = intial_counts + 1

    second_first_first = document.add_paragraph("工况参数")
    run = second_first_first.runs[0]
    run.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)  # 设置文本颜色为红色
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    run.font.size = docx.shared.Pt(size_1)  # 设置第一个运行的字体大小为15磅
    intial_counts = intial_counts + 1

    for item in work_condition_params.keys():
        second_first_first = document.add_paragraph(f"{item}:{work_condition_params[item]}")
        run = second_first_first.runs[0]
        run.font.color.rgb = RGBColor(0, 0, 0)  # 设置文本颜色为红色
        run.font.name = '宋体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
        run.font.size = docx.shared.Pt(size_1)  # 设置第一个运行的字体大小为15磅
        intial_counts = intial_counts + 1

    second_first_first = document.add_paragraph("编号参数")
    run = second_first_first.runs[0]
    run.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)  # 设置文本颜色为红色
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    run.font.size = docx.shared.Pt(size_1)  # 设置第一个运行的字体大小为15磅
    intial_counts = intial_counts + 1

    for item in order_params.keys():
        second_first_first = document.add_paragraph(f"{item}:{order_params[item]}")
        run = second_first_first.runs[0]
        run.font.color.rgb = RGBColor(0, 0, 0)  # 设置文本颜色为红色
        run.font.name = '宋体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
        run.font.size = docx.shared.Pt(size_1)  # 设置第一个运行的字体大小为15磅
        intial_counts = intial_counts + 1

    second_first_first = document.add_paragraph("结构示意")
    run = second_first_first.runs[0]
    run.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)  # 设置文本颜色为红色
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    run.font.size = docx.shared.Pt(size_1)  # 设置第一个运行的字体大小为15磅
    intial_counts = intial_counts + 1

    picture = document.add_picture(os.path.join(parent_path, 'file/basepic/3dmodel.png'),
                                   width=Inches(2.25))  # 设置图片宽度，inches（英尺）与cm（厘米）两种
    document.paragraphs[intial_counts].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 居中对齐
    intial_counts = intial_counts + 1

    # 添加标题
    title = document.add_paragraph("图1、三维模型")
    run = title.runs[0]
    run.font.color.rgb = RGBColor(0, 0, 0)  # 设置文本颜色为红色
    # title.text = "仿真分析报告"
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 居中对齐
    # 设置西体格式
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    run.font.size = docx.shared.Pt(size_1)  # 设置第一个运行的字体大小为15磅
    intial_counts = intial_counts + 1

    picture = document.add_picture(os.path.join(parent_path, 'file/basepic/section.png'),
                                   width=Inches(4.25))  # 设置图片宽度，inches（英尺）与cm（厘米）两种
    document.paragraphs[intial_counts].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 居中对齐
    intial_counts = intial_counts + 1

    # 添加标题
    title = document.add_paragraph("图2、截面示意图")
    run = title.runs[0]
    run.font.color.rgb = RGBColor(0, 0, 0)  # 设置文本颜色为红色
    # title.text = "仿真分析报告"
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 居中对齐
    # 设置西体格式
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    run.font.size = docx.shared.Pt(size_1)  # 设置第一个运行的字体大小为15磅
    intial_counts = intial_counts + 1

    second_first = document.add_heading("2.监测参数", level=3)
    run = second_first.runs[0]
    run.font.color.rgb = RGBColor(0, 0, 0)  # 设置文本颜色为红色
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    run.font.size = docx.shared.Pt(size_1)  # 设置第一个运行的字体大小为15磅
    intial_counts = intial_counts + 1

    second_first = document.add_paragraph("截面参数")
    run = second_first.runs[0]
    run.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)  # 设置文本颜色为红色
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    run.font.size = docx.shared.Pt(size_1)  # 设置第一个运行的字体大小为15磅
    intial_counts = intial_counts + 1

    picture = document.add_picture(os.path.join(parent_path, f'run/{run_dir}/xz_ave_pre.png'),
                                   width=Inches(4.25))  # 设置图片宽度，inches（英尺）与cm（厘米）两种
    document.paragraphs[intial_counts].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 居中对齐
    intial_counts = intial_counts + 1
    # 截面平均压力/温度——时间图
    # 添加标题
    title = document.add_paragraph("图3、xz截面平均压力——时间图(Y轴)")
    run = title.runs[0]
    run.font.color.rgb = RGBColor(0, 0, 0)  # 设置文本颜色为红色
    # title.text = "仿真分析报告"
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 居中对齐
    # 设置西体格式
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    run.font.size = docx.shared.Pt(size_1)  # 设置第一个运行的字体大小为15磅
    intial_counts = intial_counts + 1

    picture = document.add_picture(os.path.join(parent_path, f'run/{run_dir}/xz_ave_tem.png'),
                                   width=Inches(4.25))  # 设置图片宽度，inches（英尺）与cm（厘米）两种
    document.paragraphs[intial_counts].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 居中对齐
    # 截面平均压力/温度——时间图
    # 添加标题
    intial_counts = intial_counts + 1
    title = document.add_paragraph("图4、xz截面平均温度——时间图(Y轴)")
    run = title.runs[0]
    run.font.color.rgb = RGBColor(0, 0, 0)  # 设置文本颜色为红色
    # title.text = "仿真分析报告"
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 居中对齐
    # 设置西体格式
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    run.font.size = docx.shared.Pt(size_1)  # 设置第一个运行的字体大小为15磅
    intial_counts = intial_counts + 1

    picture = document.add_picture(os.path.join(parent_path, f'run/{run_dir}/yz_ave_tem.png'),
                                   width=Inches(4.25))  # 设置图片宽度，inches（英尺）与cm（厘米）两种
    document.paragraphs[intial_counts].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 居中对齐
    # 截面平均压力/温度——时间图
    # 添加标题
    intial_counts = intial_counts + 1
    title = document.add_paragraph("图5、yz截面平均温度——时间图(Y轴)")
    run = title.runs[0]
    run.font.color.rgb = RGBColor(0, 0, 0)  # 设置文本颜色为红色
    # title.text = "仿真分析报告"
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 居中对齐
    # 设置西体格式
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    run.font.size = docx.shared.Pt(size_1)  # 设置第一个运行的字体大小为15磅
    intial_counts = intial_counts + 1

    picture = document.add_picture(os.path.join(parent_path, f'run/{run_dir}/yz_ave_pre.png'),
                                   width=Inches(4.25))  # 设置图片宽度，inches（英尺）与cm（厘米）两种
    document.paragraphs[intial_counts].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 居中对齐
    intial_counts = intial_counts + 1
    # 截面平均压力/温度——时间图
    # 添加标题
    title = document.add_paragraph("图6、yz截面平均压力——时间图(Y轴)")
    run = title.runs[0]
    run.font.color.rgb = RGBColor(0, 0, 0)  # 设置文本颜色为红色
    # title.text = "仿真分析报告"
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 居中对齐
    # 设置西体格式
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    run.font.size = docx.shared.Pt(size_1)  # 设置第一个运行的字体大小为15磅
    intial_counts = intial_counts + 1

    second_first = document.add_paragraph("体积参数")
    run = second_first.runs[0]
    run.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)  # 设置文本颜色为红色
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    run.font.size = docx.shared.Pt(size_1)  # 设置第一个运行的字体大小为15磅
    intial_counts = intial_counts + 1

    picture = document.add_picture(os.path.join(parent_path, f'run/{run_dir}/volume_liquid_time.png'),
                                   width=Inches(4.25))  # 设置图片宽度，inches（英尺）与cm（厘米）两种
    document.paragraphs[intial_counts].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 居中对齐
    intial_counts = intial_counts + 1
    # 截面平均压力/温度——时间图
    # 添加标题
    title = document.add_paragraph("图7、液体体积——时间图(Y轴)")
    run = title.runs[0]
    run.font.color.rgb = RGBColor(0, 0, 0)  # 设置文本颜色为红色
    # title.text = "仿真分析报告"
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 居中对齐
    # 设置西体格式
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    run.font.size = docx.shared.Pt(size_1)  # 设置第一个运行的字体大小为15磅
    intial_counts = intial_counts + 1

    picture = document.add_picture(os.path.join(parent_path, f'run/{run_dir}/mass_liquid_time.png'),
                                   width=Inches(4.25))  # 设置图片宽度，inches（英尺）与cm（厘米）两种
    document.paragraphs[intial_counts].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 居中对齐
    intial_counts = intial_counts + 1
    # 截面平均压力/温度——时间图
    # 添加标题
    title = document.add_paragraph("图8、液体质量——时间图(Y轴)")
    run = title.runs[0]
    run.font.color.rgb = RGBColor(0, 0, 0)  # 设置文本颜色为红色
    # title.text = "仿真分析报告"
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 居中对齐
    # 设置西体格式
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    run.font.size = docx.shared.Pt(size_1)  # 设置第一个运行的字体大小为15磅
    intial_counts = intial_counts + 1

    picture = document.add_picture(os.path.join(parent_path, f'run/{run_dir}/volume_gas_time.png'),
                                   width=Inches(4.25))  # 设置图片宽度，inches（英尺）与cm（厘米）两种
    #
    document.paragraphs[intial_counts].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 居中对齐
    intial_counts = intial_counts + 1
    # 截面平均压力/温度——时间图
    # 添加标题
    title = document.add_paragraph("图9、气体体积——时间图(Y轴)")
    run = title.runs[0]
    run.font.color.rgb = RGBColor(0, 0, 0)  # 设置文本颜色为红色
    # title.text = "仿真分析报告"
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 居中对齐
    # 设置西体格式
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    run.font.size = docx.shared.Pt(size_1)  # 设置第一个运行的字体大小为15磅
    intial_counts = intial_counts + 1

    picture = document.add_picture(os.path.join(parent_path, f'run/{run_dir}/mass_gas_time.png'),
                                   width=Inches(4.25))  # 设置图片宽度，inches（英尺）与cm（厘米）两种
    document.paragraphs[intial_counts].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 居中对齐
    intial_counts = intial_counts + 1
    # 截面平均压力/温度——时间图
    # 添加标题
    title = document.add_paragraph("图10、气体质量——时间图(Y轴)")
    run = title.runs[0]
    run.font.color.rgb = RGBColor(0, 0, 0)  # 设置文本颜色为红色
    # title.text = "仿真分析报告"
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 居中对齐
    # 设置西体格式
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    run.font.size = docx.shared.Pt(size_1)  # 设置第一个运行的字体大小为15磅
    intial_counts = intial_counts + 1

    second_first = document.add_heading("3. 图片显示(模型)", level=3)
    run = second_first.runs[0]
    run.font.color.rgb = RGBColor(0, 0, 0)  # 设置文本颜色为红色
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    run.font.size = docx.shared.Pt(size_1)  # 设置第一个运行的字体大小为15磅
    intial_counts = intial_counts + 1

    second_first = document.add_paragraph("侧视云图")
    run = second_first.runs[0]
    run.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)  # 设置文本颜色为红色
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    run.font.size = docx.shared.Pt(size_1)  # 设置第一个运行的字体大小为15磅
    intial_counts = intial_counts + 1

    picture = document.add_picture(os.path.join(parent_path, f'run/{run_dir}/last_pre_conture.png'),
                                   width=Inches(3.25))  # 设置图片宽度，inches（英尺）与cm（厘米）两种
    document.paragraphs[intial_counts].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 居中对齐
    intial_counts = intial_counts + 1
    # 截面平均压力/温度——时间图
    # 添加标题
    title = document.add_paragraph("图11、压力云图图片—终止时刻")
    run = title.runs[0]
    run.font.color.rgb = RGBColor(0, 0, 0)  # 设置文本颜色为红色
    # title.text = "仿真分析报告"
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 居中对齐
    # 设置西体格式
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    run.font.size = docx.shared.Pt(size_1)  # 设置第一个运行的字体大小为15磅
    intial_counts = intial_counts + 1

    picture = document.add_picture(os.path.join(parent_path, f'run/{run_dir}/last_liquid_level_conture.png'),
                                   width=Inches(3.25))  # 设置图片宽度，inches（英尺）与cm（厘米）两种
    document.paragraphs[intial_counts].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 居中对齐
    # 截面平均压力/温度——时间图
    # 添加标题
    intial_counts = intial_counts + 1
    title = document.add_paragraph("图12、液位云图图片—终止时刻")
    run = title.runs[0]
    run.font.color.rgb = RGBColor(0, 0, 0)  # 设置文本颜色为红色
    # title.text = "仿真分析报告"
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 居中对齐
    # 设置西体格式
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    run.font.size = docx.shared.Pt(size_1)  # 设置第一个运行的字体大小为15磅
    intial_counts = intial_counts + 1

    pass
