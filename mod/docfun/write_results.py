from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import RGBColor
from docx.oxml.ns import qn  # 中文格式
from docx.shared import Inches, Cm
import pandas as pd
import docx

from mod.docfun.table_plot import add_three_line


def write_results(document,initla_num,pre_table:pd.DataFrame,sus_t,time_cost):
    size_0 = 16
    size_1 = 10
    size_2 = 8
    second = document.add_heading("三、结果分析", level=2)
    run = second.runs[0]
    run.font.color.rgb = RGBColor(0, 0, 0)  # 设置文本颜色为红色
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    run.font.size = docx.shared.Pt(size_0)  # 设置第一个运行的字体大小为15磅
    initla_num = initla_num+1

    title = document.add_paragraph("计算结果分析", style='List Number')
    run = title.runs[0]
    run.font.color.rgb = RGBColor(0, 0, 0)  # 设置文本颜色为红色
    # title.text = "仿真分析报告"
    # 设置西体格式
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    run.font.size = docx.shared.Pt(size_1)  # 设置第一个运行的字体大小为15磅
    initla_num = initla_num + 1

    paragraph = document.add_paragraph(f"表1、升压速率分析——截面平均压力变化速率")
    run = paragraph.runs[0]
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 居中对齐
    r = paragraph.runs[0]  # 获取第一个运行
    r.font.size = docx.shared.Pt(size_1)  # 设置第一个运行的字体大小为15磅
    initla_num = initla_num + 1

    v_cell_num = pre_table.shape[0] + 1  # 4
    h_cell_num = pre_table.shape[1] + 1  # 3

    # 添加表格
    table = document.add_table(rows=v_cell_num, cols=h_cell_num)
    initla_num = initla_num + 1

    # 写入内容
    for i in range(1, v_cell_num):
        table.cell(i, 0).text = str(pre_table.index[i - 1])
        for j in range(1, h_cell_num):
            cell = table.cell(i, j)
            cell.text = f"{pre_table.iloc[i - 1, j - 1]}"
            if i == 1:
                table.cell(0, j).text = str(pre_table.columns[j - 1])
    # 写入表头
    table = add_three_line(v_cell_num, h_cell_num, table)
    # 设置表格的样式
    table.style = 'Table Grid'
    # 保存文档

    paragraph = document.add_paragraph(f"表2、维持时间分析——最大压力（设定值）/升压速率")
    run = paragraph.runs[0]
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 居中对齐
    r = paragraph.runs[0]  # 获取第一个运行
    r.font.size = docx.shared.Pt(size_2)  # 设置第一个运行的字体大小为15磅
    # 维持时间分析——最大压力（设定值）/升压速率
    # 示例DataFrame
    initla_num = initla_num + 1

    v_cell_num = sus_t.shape[0] + 1  # 4
    h_cell_num = sus_t.shape[1] + 1  # 3

    # 添加表格
    table = document.add_table(rows=v_cell_num, cols=h_cell_num)

    # 写入内容
    for i in range(1, v_cell_num):
        table.cell(i, 0).text = str(sus_t.index[i - 1])
        for j in range(1, h_cell_num):
            cell = table.cell(i, j)
            cell.text = f"{sus_t.iloc[i - 1, j - 1]}"
            if i == 1:
                table.cell(0, j).text = str(sus_t.columns[j - 1])
    # 写入表头

    table = add_three_line(v_cell_num, h_cell_num, table)
    # 设置表格的样式
    table.style = 'Table Grid'
    initla_num = initla_num + 1

    title = document.add_paragraph("计算分析", style='List Number')
    run = title.runs[0]
    run.font.color.rgb = RGBColor(0, 0, 0)  # 设置文本颜色为红色
    # title.text = "仿真分析报告"
    # 设置西体格式
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    run.font.size = docx.shared.Pt(size_1)  # 设置第一个运行的字体大小为15磅
    initla_num = initla_num + 1


    title = document.add_paragraph(f"计算时长{time_cost}s")
    run = title.runs[0]
    run.font.color.rgb = RGBColor(0, 0, 0)  # 设置文本颜色为红色
    # title.text = "仿真分析报告"
    # 设置西体格式
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    run.font.size = docx.shared.Pt(size_1)  # 设置第一个运行的字体大小为15磅
    initla_num = initla_num + 1

    return initla_num

