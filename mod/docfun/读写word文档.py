# -*- coding: utf-8 -*-
"""
Created on 2024/1/30 14:01

@File: 读写word文档.py

@Department:

@Author: jianqiugao

@Email: 985040320@qq.com

@Describe: **加入文件描述, 要实现的功能, 注意事项等**
"""
import docx
from docx import Document
from docx.shared import Inches, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml.ns import qn  # 中文格式
from table_plot import add_three_line



# 创建一个新的Word文档对象
document = Document()

# 添加标题
title = document.add_heading("后处理报告参数内容汇总", level=1)
run = title.runs[0]
run.font.color.rgb = RGBColor(0, 0, 0)  # 设置文本颜色为红色
# title.text = "仿真分析报告"
title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 居中对齐
# 设置西体格式
run.font.name = '宋体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
run.font.size = docx.shared.Pt(16)  # 设置第一个运行的字体大小为15磅



# run.font.highlight_color = WD_COLOR_INDEX.YELLOW # 设置文本底纹颜色为黄色
first = document.add_heading("一、参考依据", level=2)
run = first.runs[0]
run.font.color.rgb = RGBColor(0, 0, 0)  # 设置文本颜色为红色
run.font.name = '宋体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
run.font.size = docx.shared.Pt(16)  # 设置第一个运行的字体大小为15磅



first_first = document.add_heading("1.技术需求", level=3)
run = first_first.runs[0]
run.font.color.rgb = RGBColor(0, 0, 0)  # 设置文本颜色为红色
run.font.name = '宋体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
run.font.size = docx.shared.Pt(10)  # 设置第一个运行的字体大小为15磅

picture = document.add_picture('../../file/pic/tech_requirement.png', width=Inches(6.25))  # 设置图片宽度，inches（英尺）与cm（厘米）两种
document.paragraphs[3].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 居中对齐

first_second = document.add_heading("2.合同约定", level=3)
run = first_second.runs[0]
run.font.color.rgb = RGBColor(0, 0, 0)  # 设置文本颜色为红色
run.font.name = '宋体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
run.font.size = docx.shared.Pt(10)  # 设置第一个运行的字体大小为15磅

first_second_first = document.add_paragraph("1).软件接口方面,基于现有商业软件ansys fluent,Ensight开发;")
run = first_second_first.runs[0]
run.font.color.rgb = RGBColor(0, 0, 0)  # 设置文本颜色为红色
run.font.name = '宋体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
run.font.size = docx.shared.Pt(8)  # 设置第一个运行的字体大小为15磅

first_second_second = document.add_paragraph("2).实现作业计算自动监控，并在分析完成后自动进行结果分析输出;")
run = first_second_second.runs[0]
run.font.color.rgb = RGBColor(0, 0, 0)  # 设置文本颜色为红色
run.font.name = '宋体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
run.font.size = docx.shared.Pt(8)  # 设置第一个运行的字体大小为15磅

first_second_third = document.add_paragraph("3).实现分析报告自动输出,报告格式为word/pdf/ppt;")
run = first_second_third.runs[0]
run.font.color.rgb = RGBColor(0, 0, 0)  # 设置文本颜色为红色
run.font.name = '宋体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
run.font.size = docx.shared.Pt(8)  # 设置第一个运行的字体大小为15磅

first_second_third = document.add_paragraph("4).报告内容要求包括但不限于仿真需求中设置的工况条件,模型设置,数据曲线对比,云图等;")
run = first_second_third.runs[0]
run.font.color.rgb = RGBColor(0, 0, 0)  # 设置文本颜色为红色
run.font.name = '宋体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
run.font.size = docx.shared.Pt(8)  # 设置第一个运行的字体大小为15磅

first_second_third = document.add_paragraph("5).开发语言采用底层计算机语言，适用于Windows操作系统;")
run = first_second_third.runs[0]
run.font.color.rgb = RGBColor(0, 0, 0)  # 设置文本颜色为红色
run.font.name = '宋体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
run.font.size = docx.shared.Pt(8)  # 设置第一个运行的字体大小为15磅

first_third = document.add_heading("3.相关标准文件", level=3)
run = first_third.runs[0]
run.font.color.rgb = RGBColor(0, 0, 0)  # 设置文本颜色为红色
run.font.name = '宋体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
run.font.size = docx.shared.Pt(10)  # 设置第一个运行的字体大小为15磅

first_third_first = document.add_paragraph("仿真结果分析及可视化应符合以下要求")
run = first_third_first.runs[0]
run.font.color.rgb = RGBColor(0, 0, 0)  # 设置文本颜色为红色
run.font.name = '宋体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
run.font.size = docx.shared.Pt(8)  # 设置第一个运行的字体大小为15磅

first_third_second = document.add_paragraph("a) 系统仿真应能够进行工时、材料定额的计算和统计，辅助进行材料及人工成本分析")
run = first_third_second.runs[0]
run.font.color.rgb = RGBColor(0, 0, 0)  # 设置文本颜色为红色
run.font.name = '宋体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
run.font.size = docx.shared.Pt(8)  # 设置第一个运行的字体大小为15磅

first_third_third = document.add_paragraph("b) 系统仿真应根据不同的系统特点、仿真要求对仿真数据进行取舍分析，得到评价结果;")
run = first_third_third.runs[0]
run.font.color.rgb = RGBColor(0, 0, 0)  # 设置文本颜色为红色
run.font.name = '宋体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
run.font.size = docx.shared.Pt(8)  # 设置第一个运行的字体大小为15磅


first_third_fourth = document.add_paragraph("c) 应能通过汇总分析加工精度、设备利用率、加工方法、加工效率等信息，综合判断NC程序是否合理，辅助进行NC程序优化，并能对优化前后评价因素进行比对，分析或辅助分析改进情况;")
run = first_third_fourth.runs[0]
run.font.color.rgb = RGBColor(0, 0, 0)  # 设置文本颜色为红色
run.font.name = '宋体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
run.font.size = docx.shared.Pt(8)  # 设置第一个运行的字体大小为15磅

first_third_fi = document.add_paragraph("d) 对于人体模型应能对工效分析结果进行分析，辅助进行人体操作改进，在进行优化改进时，应d能从姿态分析、视野分析、作业空间、工效分析等方面对比优化改进情况;")
run = first_third_fi.runs[0]
run.font.color.rgb = RGBColor(0, 0, 0)  # 设置文本颜色为红色
run.font.name = '宋体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
run.font.size = docx.shared.Pt(8)  # 设置第一个运行的字体大小为15磅


first_third_six = document.add_paragraph("e) 生产线系统仿真结果评价应包含生产线瓶颈识别和分析，包括产生瓶颈的设备位置和程度，其指标有生产线的瓶颈率、等待队长、等待时间等:")
run = first_third_six.runs[0]
run.font.color.rgb = RGBColor(0, 0, 0)  # 设置文本颜色为红色
run.font.name = '宋体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
run.font.size = docx.shared.Pt(8)  # 设置第一个运行的字体大小为15磅

first_third_seven = document.add_paragraph("f) 生产线系统仿真结果评价应包含车间产能分析，包括最大能力和负载率，其指标包括单位时间生产量、初始加工时间、设备利用率等;")
run = first_third_seven.runs[0]
run.font.color.rgb = RGBColor(0, 0, 0)  # 设置文本颜色为红色
run.font.name = '宋体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
run.font.size = docx.shared.Pt(8)  # 设置第一个运行的字体大小为15磅

first_third_seven = document.add_paragraph("f) 生产线系统仿真结果评价应包含车间产能分析，包括最大能力和负载率，其指标包括单位时间生产量、初始加工时间、设备利用率等;")
run = first_third_seven.runs[0]
run.font.color.rgb = RGBColor(0, 0, 0)  # 设置文本颜色为红色
run.font.name = '宋体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
run.font.size = docx.shared.Pt(8)  # 设置第一个运行的字体大小为15磅

first_third_seven = document.add_paragraph("g) 生产线系统仿真结果评价还应包含生产周期分析以及生产的人力需求分析;")
run = first_third_seven.runs[0]
run.font.color.rgb = RGBColor(0, 0, 0)  # 设置文本颜色为红色
run.font.name = '宋体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
run.font.size = docx.shared.Pt(8)  # 设置第一个运行的字体大小为15磅

second = document.add_heading("二、后处理框架内容", level=2)
run = second.runs[0]
run.font.color.rgb = RGBColor(0, 0, 0)  # 设置文本颜色为红色
run.font.name = '宋体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
run.font.size = docx.shared.Pt(16)  # 设置第一个运行的字体大小为15磅

second_first = document.add_heading("1.输入参数", level=3)
run = second_first.runs[0]
run.font.color.rgb = RGBColor(0, 0, 0)  # 设置文本颜色为红色
run.font.name = '宋体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
run.font.size = docx.shared.Pt(10)  # 设置第一个运行的字体大小为15磅


second_first_first = document.add_paragraph("工况参数")
run = second_first_first.runs[0]
run.bold = True
run.font.color.rgb = RGBColor(0, 0, 0)  # 设置文本颜色为红色
run.font.name = '宋体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
run.font.size = docx.shared.Pt(10)  # 设置第一个运行的字体大小为15磅


work_condition_params = {'壁面参数':None,"入口参数":None,"出口参数":None,"加速度参数":None,"液位参数":None,"时间参数":None}
for item in work_condition_params.keys():
    second_first_first = document.add_paragraph(f"{item}:{work_condition_params[item]}")
    run = second_first_first.runs[0]
    run.font.color.rgb = RGBColor(0, 0, 0)  # 设置文本颜色为红色
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    run.font.size = docx.shared.Pt(10)  # 设置第一个运行的字体大小为15磅

second_first_first = document.add_paragraph("编号参数")
run = second_first_first.runs[0]
run.bold = True
run.font.color.rgb = RGBColor(0, 0, 0)  # 设置文本颜色为红色
run.font.name = '宋体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
run.font.size = docx.shared.Pt(10)  # 设置第一个运行的字体大小为15磅


order_params = {'名称参数':None,"编号参数":None}
for item in order_params.keys():
    second_first_first = document.add_paragraph(f"{item}:{order_params[item]}")
    run = second_first_first.runs[0]
    run.font.color.rgb = RGBColor(0, 0, 0)  # 设置文本颜色为红色
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    run.font.size = docx.shared.Pt(10)  # 设置第一个运行的字体大小为15磅


second_first_first = document.add_paragraph("结构示意")
run = second_first_first.runs[0]
run.bold = True
run.font.color.rgb = RGBColor(0, 0, 0)  # 设置文本颜色为红色
run.font.name = '宋体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
run.font.size = docx.shared.Pt(10)  # 设置第一个运行的字体大小为15磅


picture = document.add_picture('../../file/pic/3dmodel.png', width=Inches(2.25))  # 设置图片宽度，inches（英尺）与cm（厘米）两种
document.paragraphs[33].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 居中对齐


# 添加标题
title = document.add_paragraph("图1、三维模型")
run = title.runs[0]
run.font.color.rgb = RGBColor(0, 0, 0)  # 设置文本颜色为红色
# title.text = "仿真分析报告"
title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 居中对齐
# 设置西体格式
run.font.name = '宋体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
run.font.size = docx.shared.Pt(10)  # 设置第一个运行的字体大小为15磅


picture = document.add_picture('../../file/pic/section.png', width=Inches(4.25))  # 设置图片宽度，inches（英尺）与cm（厘米）两种
document.paragraphs[35].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 居中对齐


# 添加标题
title = document.add_paragraph("图2、截面示意图")
run = title.runs[0]
run.font.color.rgb = RGBColor(0, 0, 0)  # 设置文本颜色为红色
# title.text = "仿真分析报告"
title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 居中对齐
# 设置西体格式
run.font.name = '宋体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
run.font.size = docx.shared.Pt(10)  # 设置第一个运行的字体大小为15磅


second_first = document.add_heading("2.监测参数", level=3)
run = second_first.runs[0]
run.font.color.rgb = RGBColor(0, 0, 0)  # 设置文本颜色为红色
run.font.name = '宋体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
run.font.size = docx.shared.Pt(10)  # 设置第一个运行的字体大小为15磅

second_first = document.add_paragraph("截面参数")
run = second_first.runs[0]
run.bold = True
run.font.color.rgb = RGBColor(0, 0, 0)  # 设置文本颜色为红色
run.font.name = '宋体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
run.font.size = docx.shared.Pt(10)  # 设置第一个运行的字体大小为15磅


picture = document.add_picture('../../run/20240201_215030/xz_ave_pre.png', width=Inches(4.25))  # 设置图片宽度，inches（英尺）与cm（厘米）两种
document.paragraphs[39].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 居中对齐
# 截面平均压力/温度——时间图
# 添加标题
title = document.add_paragraph("图3、xz截面平均压力——时间图(Y轴)")
run = title.runs[0]
run.font.color.rgb = RGBColor(0, 0, 0)  # 设置文本颜色为红色
# title.text = "仿真分析报告"
title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 居中对齐
# 设置西体格式
run.font.name = '宋体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
run.font.size = docx.shared.Pt(10)  # 设置第一个运行的字体大小为15磅


picture = document.add_picture('../../run/20240201_215030/xz_ave_tem.png', width=Inches(4.25))  # 设置图片宽度，inches（英尺）与cm（厘米）两种
document.paragraphs[41].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 居中对齐
# 截面平均压力/温度——时间图
# 添加标题
title = document.add_paragraph("图4、xz截面平均温度——时间图(Y轴)")
run = title.runs[0]
run.font.color.rgb = RGBColor(0, 0, 0)  # 设置文本颜色为红色
# title.text = "仿真分析报告"
title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 居中对齐
# 设置西体格式
run.font.name = '宋体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
run.font.size = docx.shared.Pt(10)  # 设置第一个运行的字体大小为15磅

picture = document.add_picture('../../run/20240201_215030/yz_ave_tem.png', width=Inches(4.25))  # 设置图片宽度，inches（英尺）与cm（厘米）两种
document.paragraphs[43].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 居中对齐
# 截面平均压力/温度——时间图
# 添加标题
title = document.add_paragraph("图5、yz截面平均温度——时间图(Y轴)")
run = title.runs[0]
run.font.color.rgb = RGBColor(0, 0, 0)  # 设置文本颜色为红色
# title.text = "仿真分析报告"
title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 居中对齐
# 设置西体格式
run.font.name = '宋体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
run.font.size = docx.shared.Pt(10)  # 设置第一个运行的字体大小为15磅

picture = document.add_picture('../../run/20240201_215030/yz_ave_pre.png', width=Inches(4.25))  # 设置图片宽度，inches（英尺）与cm（厘米）两种
document.paragraphs[45].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 居中对齐
# 截面平均压力/温度——时间图
# 添加标题
title = document.add_paragraph("图6、yz截面平均压力——时间图(Y轴)")
run = title.runs[0]
run.font.color.rgb = RGBColor(0, 0, 0)  # 设置文本颜色为红色
# title.text = "仿真分析报告"
title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 居中对齐
# 设置西体格式
run.font.name = '宋体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
run.font.size = docx.shared.Pt(10)  # 设置第一个运行的字体大小为15磅


second_first = document.add_paragraph("体积参数")
run = second_first.runs[0]
run.bold = True
run.font.color.rgb = RGBColor(0, 0, 0)  # 设置文本颜色为红色
run.font.name = '宋体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
run.font.size = docx.shared.Pt(10)  # 设置第一个运行的字体大小为15磅


picture = document.add_picture('../../run/20240201_215030/volume_liquid_time.png', width=Inches(4.25))  # 设置图片宽度，inches（英尺）与cm（厘米）两种
document.paragraphs[48].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 居中对齐
# 截面平均压力/温度——时间图
# 添加标题
title = document.add_paragraph("图7、液体体积——时间图(Y轴)")
run = title.runs[0]
run.font.color.rgb = RGBColor(0, 0, 0)  # 设置文本颜色为红色
# title.text = "仿真分析报告"
title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 居中对齐
# 设置西体格式
run.font.name = '宋体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
run.font.size = docx.shared.Pt(10)  # 设置第一个运行的字体大小为15磅


picture = document.add_picture('../../run/20240201_215030/mass_liquid_time.png', width=Inches(4.25))  # 设置图片宽度，inches（英尺）与cm（厘米）两种
document.paragraphs[50].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 居中对齐
# 截面平均压力/温度——时间图
# 添加标题
title = document.add_paragraph("图8、液体质量——时间图(Y轴)")
run = title.runs[0]
run.font.color.rgb = RGBColor(0, 0, 0)  # 设置文本颜色为红色
# title.text = "仿真分析报告"
title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 居中对齐
# 设置西体格式
run.font.name = '宋体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
run.font.size = docx.shared.Pt(10)  # 设置第一个运行的字体大小为15磅


picture = document.add_picture('../../run/20240201_215030/volume_gas_time.png', width=Inches(4.25))  # 设置图片宽度，inches（英尺）与cm（厘米）两种
document.paragraphs[52].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 居中对齐
# 截面平均压力/温度——时间图
# 添加标题
title = document.add_paragraph("图9、气体体积——时间图(Y轴)")
run = title.runs[0]
run.font.color.rgb = RGBColor(0, 0, 0)  # 设置文本颜色为红色
# title.text = "仿真分析报告"
title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 居中对齐
# 设置西体格式
run.font.name = '宋体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
run.font.size = docx.shared.Pt(10)  # 设置第一个运行的字体大小为15磅


picture = document.add_picture('../../run/20240201_215030/mass_gas_time.png', width=Inches(4.25))  # 设置图片宽度，inches（英尺）与cm（厘米）两种
document.paragraphs[54].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 居中对齐
# 截面平均压力/温度——时间图
# 添加标题
title = document.add_paragraph("图10、气体质量——时间图(Y轴)")
run = title.runs[0]
run.font.color.rgb = RGBColor(0, 0, 0)  # 设置文本颜色为红色
# title.text = "仿真分析报告"
title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 居中对齐
# 设置西体格式
run.font.name = '宋体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
run.font.size = docx.shared.Pt(10)  # 设置第一个运行的字体大小为15磅


second_first = document.add_heading("3. 图片显示(模型)", level=3)
run = second_first.runs[0]
run.font.color.rgb = RGBColor(0, 0, 0)  # 设置文本颜色为红色
run.font.name = '宋体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
run.font.size = docx.shared.Pt(10)  # 设置第一个运行的字体大小为15磅

second_first = document.add_paragraph("侧视云图")
run = second_first.runs[0]
run.bold = True
run.font.color.rgb = RGBColor(0, 0, 0)  # 设置文本颜色为红色
run.font.name = '宋体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
run.font.size = docx.shared.Pt(10)  # 设置第一个运行的字体大小为15磅


picture = document.add_picture('../../run/20240201_215030/last_pre_conture.png', width=Inches(3.25))  # 设置图片宽度，inches（英尺）与cm（厘米）两种
document.paragraphs[58].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 居中对齐
# 截面平均压力/温度——时间图
# 添加标题
title = document.add_paragraph("图11、压力云图图片—终止时刻")
run = title.runs[0]
run.font.color.rgb = RGBColor(0, 0, 0)  # 设置文本颜色为红色
# title.text = "仿真分析报告"
title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 居中对齐
# 设置西体格式
run.font.name = '宋体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
run.font.size = docx.shared.Pt(10)  # 设置第一个运行的字体大小为15磅


picture = document.add_picture('../../run/20240201_215030/last_liquid_level_conture.png', width=Inches(3.25))  # 设置图片宽度，inches（英尺）与cm（厘米）两种
document.paragraphs[60].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 居中对齐
# 截面平均压力/温度——时间图
# 添加标题
title = document.add_paragraph("图12、液位云图图片—终止时刻")
run = title.runs[0]
run.font.color.rgb = RGBColor(0, 0, 0)  # 设置文本颜色为红色
# title.text = "仿真分析报告"
title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 居中对齐
# 设置西体格式
run.font.name = '宋体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
run.font.size = docx.shared.Pt(10)  # 设置第一个运行的字体大小为15磅


second = document.add_heading("三、结果分析", level=2)
run = second.runs[0]
run.font.color.rgb = RGBColor(0, 0, 0)  # 设置文本颜色为红色
run.font.name = '宋体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
run.font.size = docx.shared.Pt(16)  # 设置第一个运行的字体大小为15磅

# data = "123*456"
# 添加段落
# paragraph = document.add_paragraph(f"升压速率分析——截面平均压力变化速率")
# run = paragraph.runs[0]
# run.font.name = '宋体'
# run.element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

title = document.add_paragraph("计算结果分析",style='List Number')
run = title.runs[0]
run.font.color.rgb = RGBColor(0, 0, 0)  # 设置文本颜色为红色
# title.text = "仿真分析报告"
# 设置西体格式
run.font.name = '宋体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
run.font.size = docx.shared.Pt(10)  # 设置第一个运行的字体大小为15磅


paragraph = document.add_paragraph(f"表1、升压速率分析——截面平均压力变化速率")
run = paragraph.runs[0]
run.font.name = '宋体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 居中对齐
r = paragraph.runs[0] # 获取第一个运行
r.font.size = docx.shared.Pt(8)  # 设置第一个运行的字体大小为15磅
import pandas as pd
# 示例DataFrame
df = pd.DataFrame({'截面平均升压速率Pa/s': [1, 2, 3,5]},index=[f'截面1','截面2','截面3','截面4'])


v_cell_num = df.shape[0]+1 # 4
h_cell_num = df.shape[1]+1 # 3

# 添加表格
table = document.add_table(rows=v_cell_num, cols=h_cell_num)

# 写入内容
for i in range(1,v_cell_num):
    table.cell(i, 0).text=str(df.index[i-1])
    for j in range(1,h_cell_num):
        cell = table.cell(i, j)
        cell.text = f"{df.iloc[i-1,j-1]}"
        if i==1:
            table.cell(0, j).text = str(df.columns[j - 1])
# 写入表头

table = add_three_line(v_cell_num,h_cell_num,table)
# 设置表格的样式
table.style = 'Table Grid'
# 保存文档

paragraph = document.add_paragraph(f"表2、维持时间分析——最大压力（设定值）/升压速率")
run = paragraph.runs[0]
run.font.name = '宋体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 居中对齐
r = paragraph.runs[0] # 获取第一个运行
r.font.size = docx.shared.Pt(8)  # 设置第一个运行的字体大小为15磅
# 维持时间分析——最大压力（设定值）/升压速率
# 示例DataFrame
df = pd.DataFrame({'维持时间t': [1, 2, 3,5],'初始时刻截面平均压力Pa': [1, 2, 3,5]},index=[f'截面1','截面2','截面3','截面4'])


v_cell_num = df.shape[0]+1 # 4
h_cell_num = df.shape[1]+1 # 3

# 添加表格
table = document.add_table(rows=v_cell_num, cols=h_cell_num)

# 写入内容
for i in range(1,v_cell_num):
    table.cell(i, 0).text=str(df.index[i-1])
    for j in range(1,h_cell_num):
        cell = table.cell(i, j)
        cell.text = f"{df.iloc[i-1,j-1]}"
        if i==1:
            table.cell(0, j).text = str(df.columns[j - 1])
# 写入表头

table = add_three_line(v_cell_num,h_cell_num,table)
# 设置表格的样式
table.style = 'Table Grid'

title = document.add_paragraph("计算分析",style='List Number')
run = title.runs[0]
run.font.color.rgb = RGBColor(0, 0, 0)  # 设置文本颜色为红色
# title.text = "仿真分析报告"
# 设置西体格式
run.font.name = '宋体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
run.font.size = docx.shared.Pt(10)  # 设置第一个运行的字体大小为15磅

time_cost =12
title = document.add_paragraph(f"计算时长{time_cost}s")
run = title.runs[0]
run.font.color.rgb = RGBColor(0, 0, 0)  # 设置文本颜色为红色
# title.text = "仿真分析报告"
# 设置西体格式
run.font.name = '宋体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
run.font.size = docx.shared.Pt(10)  # 设置第一个运行的字体大小为15磅

second = document.add_heading("四、报告总结", level=2)
run = second.runs[0]
run.font.color.rgb = RGBColor(0, 0, 0)  # 设置文本颜色为红色
run.font.name = '宋体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
run.font.size = docx.shared.Pt(16)  # 设置第一个运行的字体大小为15磅
name = "20240210_215030"

title = document.add_paragraph(f"报告名称:{name}")
run = title.runs[0]
run.font.color.rgb = RGBColor(0, 0, 0)  # 设置文本颜色为红色
# title.text = "仿真分析报告"
# 设置西体格式
run.font.name = '宋体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
run.font.size = docx.shared.Pt(10)  # 设置第一个运行的字体大小为15磅

sumary = "计算顺利完成，储罐。。。。。。。。。。。。。"
# 计算结论
title = document.add_paragraph(f"计算结论:{sumary}")
run = title.runs[0]
run.font.color.rgb = RGBColor(0, 0, 0)  # 设置文本颜色为红色
# title.text = "仿真分析报告"
# 设置西体格式
run.font.name = '宋体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
run.font.size = docx.shared.Pt(10)  # 设置第一个运行的字体大小为15磅

# 编写
editor = 'computer auto writer'
title = document.add_paragraph(f"编写:{editor}")
run = title.runs[0]
run.font.color.rgb = RGBColor(0, 0, 0)  # 设置文本颜色为红色
# title.text = "仿真分析报告"
# 设置西体格式
run.font.name = '宋体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
run.font.size = docx.shared.Pt(10)  # 设置第一个运行的字体大小为15磅

# 审核
checker = '高建秋'
title = document.add_paragraph(f"编写:{checker}")
run = title.runs[0]
run.font.color.rgb = RGBColor(0, 0, 0)  # 设置文本颜色为红色
# title.text = "仿真分析报告"
# 设置西体格式
run.font.name = '宋体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
run.font.size = docx.shared.Pt(10)  # 设置第一个运行的字体大小为15磅

# 审核
author = '高建秋'
title = document.add_paragraph(f"审核:{author}")
run = title.runs[0]
run.font.color.rgb = RGBColor(0, 0, 0)  # 设置文本颜色为红色
# title.text = "仿真分析报告"
# 设置西体格式
run.font.name = '宋体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
run.font.size = docx.shared.Pt(10)  # 设置第一个运行的字体大小为15磅

import datetime
# 审核
data = datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')
title = document.add_paragraph(f"日期:{data}")
run = title.runs[0]
run.font.color.rgb = RGBColor(0, 0, 0)  # 设置文本颜色为红色
# title.text = "仿真分析报告"
# 设置西体格式
run.font.name = '宋体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
run.font.size = docx.shared.Pt(10)  # 设置第一个运行的字体大小为15磅

document.save('example1.docx')
